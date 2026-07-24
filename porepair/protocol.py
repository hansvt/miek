"""WI-7 — reproducibility / methods protocol as a .docx, populated from the actual run
(not hard-coded), meant to be rewritten by the user for publication."""
import os
from docx import Document
from docx.shared import Pt


def _p(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead + " "); r.bold = True
    p.add_run(text)
    return p


def build(out_dir, results, meta):
    r = results
    t = r["transform"]
    cal = r["calibration_um_per_px"]
    mrd = r.get("match_radius_detail", {})
    ipd = r["interpore_distance_um"]
    oct_name = os.path.basename(r["images"]["oct"])
    imm_name = os.path.basename(r["images"]["imm"])

    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)
    doc.add_heading("Methods — OCT ↔ immunolabel fingerprint pore analysis", level=1)
    _p(doc, "This methods section is generated from the actual analysis run; edit freely for "
            "publication. All numeric values reflect this run's configuration and results.").italic = True

    doc.add_heading("Image sources and calibration", level=2)
    _p(doc, f"Two images of the same fingerprint were analysed: an OCT en-face image "
            f"({oct_name}) and an immunolabel fluorescence image ({imm_name}). Physical scale "
            f"was taken from the OCT calibration: the OCT field is "
            f"{cal['oct_mm'][0]:.0f} × {cal['oct_mm'][1]:.0f} mm, giving {cal['x']:.2f} µm/px "
            f"(x) and {cal['y']:.2f} µm/px (y). All distances are reported in micrometres using "
            f"this calibration; the immunolabel scale follows from the registration transform.")

    doc.add_heading("Immunolabel pore detection", level=2)
    idet = r.get("imm_detection") or {}
    p = idet.get("params") or {}
    opt = p.get("optimize", {})
    rej = idet.get("rejected") or {}
    rej_txt = ", ".join(f"{k}: {v}" for k, v in rej.items()) if rej else "none"
    region = ("a user-selected region mask" if idet.get("region_mask") else
              "the print body (largest connected component after heavy blur/threshold)")
    _p(doc, f"Because the immunolabel signal is the excreted material around each pore rather than "
            f"the pore itself, pores appear as small blobs. Detection used the raw "
            f"{meta.get('imm_channel','red')} channel and was restricted to {region}. The image was "
            f"first optimised ({opt.get('method','background subtraction + contrast stretch')}; "
            f"background sigma {opt.get('bg_sigma','~2x pore diameter')}, percentile clip "
            f"{opt.get('clip_pct','[1, 99.5]')}) to flatten ridge/illumination and balance black/white; "
            f"CLAHE was used only for display, not detection. This optimisation is a documented "
            f"provisional step, intended to be replaced by the laboratory's MATLAB white/black-balance "
            f"procedure. The optimised image was binarised (Otsu) and connected components analysed. "
            f"Thresholds were scale-aware, derived from the estimated pore diameter "
            f"({p.get('pore_diam_px','?')} px): components were kept as pores when their area fell in "
            f"{p.get('min_area','?')}–{p.get('max_area','?')} px², circularity ≥ "
            f"{p.get('circularity_thresh','?')}, solidity ≥ {p.get('solidity_thresh','?')} and "
            f"eccentricity ≤ {p.get('max_eccentricity','?')}; fused blobs were handled by "
            f"'{p.get('merged_blobs','split')}'. Each pore's centroid was its position and its "
            f"equivalent radius recorded. Rejected artefacts (with reason): {rej_txt}. "
            f"A top-hat point-maximum detector was also run for comparison "
            f"({idet.get('detector','region')} was the primary detector).")

    doc.add_heading("OCT pore detection", level=2)
    _p(doc, "OCT pores (bright dots) were detected with a white top-hat followed by local-maxima "
            "detection, with the threshold set adaptively at the knee of the peak-count-versus-"
            "threshold curve. Detection was restricted to valid tissue (excluding the black "
            "scan-artefact bands).")

    doc.add_heading("Registration", level=2)
    reg = (f" A second, local refinement then used matched pore correspondences "
           f"({t.get('n_pore_pairs_used',0)} pairs, landmarks weighted higher); the landmark "
           f"residual changed from {t.get('residual_landmarks_only_px')} to "
           f"{t.get('residual_after_refine_px')} px."
           if t.get("refined_with_pores") else "")
    _p(doc, f"The OCT image was registered to the immunolabel image with a {t['kind']} transform "
            f"estimated from {t['n_points']} manually placed corresponding landmarks (ridge "
            f"endings, bifurcations, core/delta and distinct pores). The transform had scale "
            f"{t['scale']:.3f}, rotation {t['rotation_deg']:.1f}° and axis angle "
            f"{t['axis_angle_deg']:.1f}° (90° = no shear), with a mean landmark residual of "
            f"{t['residual_mean_px']:.1f} px (max {t['residual_max_px']:.1f}).{reg} Fully automatic "
            f"registration was evaluated but rejected: the quasi-regular pore lattice and slowly "
            f"varying ridge orientation make it ambiguous, so manual landmarks were used.")

    doc.add_heading("Pore counting and matching", level=2)
    _p(doc, f"The area of interest (AOI) was the full valid OCT image area ({r['roi_mm2']:.1f} mm²), "
            f"mapped into the immunolabel frame via the transform so the same area was used in both "
            f"images; {r['roi_covered_by_print_pct']}% of it lies within the labelled print. Pores "
            f"were counted in both modalities within the AOI and matched with mutual nearest-neighbour "
            f"matching. The match radius was {r['match_radius_px']} px "
            f"= {mrd.get('base_frac_x_NN')} px ({mrd.get('match_frac')}× the median pore spacing) "
            f"+ {mrd.get('margin_k_x_region')} px ({mrd.get('margin_k')}× the median pore-region "
            f"radius {mrd.get('median_region_radius_px')} px), the latter allowing for the immunolabel "
            f"halo. Pores were classified as shared, OCT-only or immuno-only.")

    doc.add_heading("Inter-pore distance and nearest-neighbour geometry", level=2)
    def med(d):
        return f"{d['median']:.0f} µm" if d else "n/a"
    _p(doc, f"Inter-pore distance was computed as the nearest-neighbour distance in micrometres. "
            f"For the matched pores it was {med(ipd.get('matched_oct'))} (OCT) and "
            f"{med(ipd.get('matched_imm'))} (immuno); across all AOI pores (a denser estimate) it "
            f"was {med(ipd.get('all_oct'))} (OCT) and {med(ipd.get('all_imm'))} (immuno). For every "
            f"pore a nearest-neighbour table records ID, coordinates, neighbour ID, distance (µm) "
            f"and angle ({r.get('nn_angle_convention','')}).")

    doc.add_heading("Software and outputs", level=2)
    _p(doc, "Analysis used the porepair pipeline (OpenCV, scikit-image, SciPy, NumPy). Outputs: "
            "report.html, results.xlsx, results.json, overlay images, matching map, numbered pore "
            "maps and nearest-neighbour tables, and this protocol.")

    path = os.path.join(out_dir, "protocol.docx")
    doc.save(path)
    return path
